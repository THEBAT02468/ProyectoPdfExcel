from docxtpl import DocxTemplate
from docx.shared import Pt
from datetime import datetime
import pandas as pd
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Carga de la plantilla y los datos
template_path = './excel_to_word/schema_excel.docx'
doc = DocxTemplate(template_path)

name = "Santiago"
numero = "123456789"
correo = "santiago@gmail.com"
direccion = "Calle"
fecha = datetime.today().strftime("%d %b, %Y")

# Contexto común para todos los documentos
context = {
    "name": name,
    "numero": numero,
    "correo": correo,
    "direccion": direccion,
    "fecha": fecha
}

# Cargar el archivo CSV
df = pd.read_csv('./excel_to_word/schema.csv', delimiter=';', encoding='UTF-8')

# Función para dar formato al documento Word
def apply_styles(doc):
    # Añadir un título centrado
    title = doc.add_paragraph()
    run = title.add_run("Contrato de Compra")
    run.bold = True
    run.font.size = Pt(16)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    
    # Configura la fuente y tamaño del texto
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    #Añadir Imagen del contrato
    image_path = './excel_to_word/contrato.png'
    #Que se centralice en el documento
    doc.add_picture(image_path, width=Pt(200), align=WD_PARAGRAPH_ALIGNMENT.CENTER)

    # Añadir una línea horizontal (separador)
    doc.add_paragraph()
    hr = doc.add_paragraph()
    hr_format = hr.add_run("─" * 50)
    hr_format.font.size = Pt(8)
    hr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Iterar sobre cada fila del DataFrame y crear un documento Word para cada una
for index, row in df.iterrows():
    # Crear el contexto específico del documento actual
    cuerpo = {
        "otro_nombre": row['name'],
        "otro_numero": row['numero'],
        "otro_correo": row['correo'],
        "otro_direccion": row['direccion'],
        "otro_fecha": fecha,
    }
    cuerpo.update(context)
    
    # Renderizar el documento con el contexto
    doc.render(cuerpo)

    doc.paragraphs.clear()  # Limpia el contenido del documento actual
    
    # Aplicar estilos personalizados al documento
    apply_styles(doc)
    
    # Guardar el documento generado con el nombre específico
    output_path = f"./excel_to_word/outputs/dato_{row['name']}.docx"
    doc.save(output_path)
    print(f"Documento generado con éxito: {output_path}")
